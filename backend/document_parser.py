"""
Document Parser Module
Extracts raw text from PDF and DOCX files for the Nexus-Prime engine.
"""

import io
from typing import Optional


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract text from a PDF file."""
    try:
        from PyPDF2 import PdfReader
        reader = PdfReader(io.BytesIO(file_bytes))
        text_parts = []
        for page in reader.pages:
            page_text = page.extract_text()
            if page_text:
                text_parts.append(page_text)
        return "\n".join(text_parts)
    except Exception as e:
        return f"[PDF_PARSE_ERROR]: {str(e)}"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract text from a DOCX file."""
    try:
        from docx import Document
        doc = Document(io.BytesIO(file_bytes))
        text_parts = []
        for paragraph in doc.paragraphs:
            if paragraph.text.strip():
                text_parts.append(paragraph.text)
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_data:
                    text_parts.append(" | ".join(row_data))
        return "\n".join(text_parts)
    except Exception as e:
        return f"[DOCX_PARSE_ERROR]: {str(e)}"


def extract_text(file_bytes: bytes, filename: str) -> Optional[dict]:
    """Route to the correct parser based on file extension."""
    lower = filename.lower()
    if lower.endswith(".pdf"):
        return {"type": "text", "content": extract_text_from_pdf(file_bytes)}
    elif lower.endswith(".docx") or lower.endswith(".doc"):
        return {"type": "text", "content": extract_text_from_docx(file_bytes)}
    elif lower.endswith(".txt"):
        return {"type": "text", "content": file_bytes.decode("utf-8", errors="replace")}
    elif lower.endswith(".csv"):
        import pandas as pd
        try:
            df = pd.read_csv(io.BytesIO(file_bytes))
            return {"type": "text", "content": df.to_markdown(index=False)}
        except Exception as e:
            return {"type": "text", "content": f"[CSV_PARSE_ERROR]: {str(e)}"}
    elif lower.endswith(".xlsx") or lower.endswith(".xls"):
        import pandas as pd
        try:
            df = pd.read_excel(io.BytesIO(file_bytes))
            return {"type": "text", "content": df.to_markdown(index=False)}
        except Exception as e:
            return {"type": "text", "content": f"[EXCEL_PARSE_ERROR]: {str(e)}"}
    elif lower.endswith(".png"):
        return {"type": "image", "mime_type": "image/png", "data": file_bytes}
    elif lower.endswith(".jpg") or lower.endswith(".jpeg"):
        return {"type": "image", "mime_type": "image/jpeg", "data": file_bytes}
    else:
        return None
